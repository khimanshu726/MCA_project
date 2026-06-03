from __future__ import annotations

import math
import os
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT_DOCX = ROOT / "ONLINE_PRINT_SHOP_MANAGEMENT_SYSTEM_MCA_REPORT.docx"
ASSETS_DIR = ROOT / "assets"
ASSETS_DIR.mkdir(parents=True, exist_ok=True)


STUDENT_NAME = "Himanshu Babu"
REG_NO = "24213110019"
SESSION = "July 2026"
COURSE = "Master of Computer Applications (MCA)"
PROJECT_TITLE = "ONLINE PRINT SHOP MANAGEMENT SYSTEM"
BUSINESS_NAME = "Elite Impressions"


def _set_run_font(run, *, name="Times New Roman", size_pt: Optional[float] = None, bold: Optional[bool] = None):
    font = run.font
    font.name = name
    # Also set East Asia name so Word sticks to Times New Roman.
    rfonts = run._element.rPr.rFonts
    rfonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold
    return run


def _set_paragraph(paragraph, *, align=None, space_before=None, space_after=None, line_spacing=None, keep_with_next=None):
    fmt = paragraph.paragraph_format
    if align is not None:
        paragraph.alignment = align
    if space_before is not None:
        fmt.space_before = Pt(space_before)
    if space_after is not None:
        fmt.space_after = Pt(space_after)
    if line_spacing is not None:
        fmt.line_spacing = line_spacing
    if keep_with_next is not None:
        fmt.keep_with_next = keep_with_next
    return paragraph


def _add_field(paragraph, field_code: str):
    # Insert a Word field (e.g., PAGE).
    run = paragraph.add_run()
    r = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r.append(fld_sep)

    text = OxmlElement("w:t")
    text.text = "1"
    r.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.append(fld_end)

    return paragraph


def _set_section_pagenum(section, *, start: int, fmt: str):
    # fmt: "roman" or "decimal"
    # OpenXML expects enum tokens like "lowerRoman" and "decimal".
    if fmt == "roman":
        fmt = "lowerRoman"
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:start"), str(start))
    pgNumType.set(qn("w:fmt"), fmt)


def _apply_page_setup(section):
    # Match a classic university-report look: A4, generous margins.
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)


def _footer_center_page_number(section, *, font_size=11):
    footer = section.footer
    # Ensure this section has an explicit footer part, otherwise some renderers
    # (and some PDF export pipelines) may not show the inherited footer.
    try:
        footer.is_linked_to_previous = False
    except Exception:
        pass
    # Ensure there is at least one paragraph.
    if not footer.paragraphs:
        p = footer.add_paragraph()
    else:
        p = footer.paragraphs[0]
        p.clear()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_field(p, " PAGE ")
    for r in p.runs:
        _set_run_font(r, size_pt=font_size)


def _doc_defaults(doc: Document):
    # Normal style defaults
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def _heading_center(doc: Document, text: str, *, size=14, bold=True, uppercase=False, space_before=12, space_after=12):
    p = doc.add_paragraph()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=space_before, space_after=space_after, line_spacing=1.5)
    run = p.add_run(text.upper() if uppercase else text)
    _set_run_font(run, size_pt=size, bold=bold)
    return p


def _heading_left(doc: Document, text: str, *, size=12, bold=True, space_before=10, space_after=6):
    p = doc.add_paragraph()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=space_before, space_after=space_after, line_spacing=1.5, keep_with_next=True)
    run = p.add_run(text)
    _set_run_font(run, size_pt=size, bold=bold)
    return p


def _body_para(doc: Document, text: str, *, space_before=0, space_after=8):
    p = doc.add_paragraph()
    _set_paragraph(
        p,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before=space_before,
        space_after=space_after,
        line_spacing=1.5,
    )
    run = p.add_run(text)
    _set_run_font(run, size_pt=12, bold=False)
    return p


def _bullets(doc: Document, items: List[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, line_spacing=1.5)
        run = p.add_run(item)
        _set_run_font(run, size_pt=12)


def _add_university_header_block(doc: Document, *, title_line: str):
    # Recreates the "institution header" feel of the sample, without hardcoding a real university.
    p = doc.add_paragraph()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, line_spacing=1.15)
    r = p.add_run("UNIVERSITY / INSTITUTE NAME")
    _set_run_font(r, size_pt=26, bold=True)
    r.font.color.rgb = RGBColor(55, 71, 79)

    p = doc.add_paragraph()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, line_spacing=1.15)
    r = p.add_run("SCHOOL OF COMPUTING")
    _set_run_font(r, size_pt=14, bold=True)

    p = doc.add_paragraph()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, line_spacing=1.15)
    r = p.add_run("CITY, STATE, INDIA - PINCODE")
    _set_run_font(r, size_pt=12, bold=True)

    _heading_center(doc, title_line, size=16, bold=True, uppercase=True, space_before=6, space_after=14)


def _add_cover_page(doc: Document):
    _add_university_header_block(doc, title_line=PROJECT_TITLE)

    _body_para(
        doc,
        'Report submitted in partial fulfillment of the requirements for the award of the degree of',
        space_before=10,
        space_after=10,
    )

    p = doc.add_paragraph()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16, line_spacing=1.15)
    r = p.add_run(COURSE)
    _set_run_font(r, size_pt=14, bold=True)
    r.font.color.rgb = RGBColor(187, 0, 0)  # subtle red like sample

    p = doc.add_paragraph()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, line_spacing=1.15)
    r = p.add_run("Submitted by")
    _set_run_font(r, size_pt=13, bold=True)

    p = doc.add_paragraph()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, line_spacing=1.15)
    r = p.add_run(STUDENT_NAME)
    _set_run_font(r, size_pt=13, bold=True)

    p = doc.add_paragraph()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14, line_spacing=1.15)
    r = p.add_run(f"(Reg. No.: {REG_NO})")
    _set_run_font(r, size_pt=12, bold=True)

    p = doc.add_paragraph()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, line_spacing=1.15)
    r = p.add_run(SESSION)
    _set_run_font(r, size_pt=12, bold=True)


def _add_bonafide(doc: Document):
    _add_university_header_block(doc, title_line="Bonafide Certificate")

    _body_para(
        doc,
        f'This is to certify that the project report titled "{PROJECT_TITLE}" submitted in fulfillment of the '
        f"requirements for the award of the degree of {COURSE} is a bona-fide record of the work done by "
        f"{STUDENT_NAME} (Reg. No. {REG_NO}) during the academic session {SESSION}. This report has not formed the basis "
        f"for the award of any degree, diploma, associateship, fellowship or other similar title to any candidate of any University.",
    )

    doc.add_paragraph()

    # Signature placeholders (must stay blank as per instructions)
    _heading_left(doc, "Signature of Project Supervisor: ______________________________", size=12, bold=True, space_before=18, space_after=6)
    _heading_left(doc, "Name with Affiliation: ______________________________", size=12, bold=True, space_before=6, space_after=6)
    _heading_left(doc, "Date: ______________________________", size=12, bold=True, space_before=6, space_after=10)
    _heading_left(doc, "Project Viva voce held on: ______________________________", size=12, bold=False, space_before=10, space_after=14)

    p = doc.add_paragraph()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=16, space_after=0)
    run = p.add_run("Examiner 1: ______________________________")
    _set_run_font(run, size_pt=12, bold=True)

    p = doc.add_paragraph()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=10, space_after=0)
    run = p.add_run("Examiner 2: ______________________________")
    _set_run_font(run, size_pt=12, bold=True)

    p = doc.add_paragraph()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=10, space_after=0)
    run = p.add_run("Seal: ______________________________")
    _set_run_font(run, size_pt=12, bold=True)


def _add_declaration(doc: Document):
    _add_university_header_block(doc, title_line="Declaration")
    _body_para(
        doc,
        f'I declare that the project report titled "{PROJECT_TITLE}" submitted by me is an original work carried out by me. '
        "The work is original and wherever I have used materials from other sources, I have given due credit and cited them "
        "in the text of the report. This report has not formed the basis for the award of any degree, diploma, associateship, "
        "fellowship or other similar title to any candidate of any University.",
    )

    doc.add_paragraph()
    _heading_left(doc, f"Candidate Name: {STUDENT_NAME}", size=12, bold=True, space_before=10, space_after=4)
    _heading_left(doc, f"Registration Number: {REG_NO}", size=12, bold=True, space_before=0, space_after=4)
    _heading_left(doc, f"Session: {SESSION}", size=12, bold=True, space_before=0, space_after=14)
    _heading_left(doc, "Signature of the Candidate: ______________________________", size=12, bold=True, space_before=16, space_after=10)
    _heading_left(doc, "Date: ______________________________", size=12, bold=True, space_before=6, space_after=0)


def _add_acknowledgement(doc: Document):
    _heading_center(doc, "Acknowledgement", size=14, bold=True, uppercase=False, space_before=14, space_after=12)
    _body_para(
        doc,
        "I express my sincere gratitude to the Institution for providing the necessary infrastructure and support to carry out "
        "this project work as a part of the MCA curriculum.",
    )
    _body_para(
        doc,
        "I am grateful to the faculty members of the Department for their continuous guidance and encouragement throughout the "
        "duration of this project.",
    )
    _body_para(
        doc,
        "I would also like to thank my project guide for the valuable suggestions, reviews, and direction that helped in completing "
        "this project successfully.",
    )
    _body_para(
        doc,
        "I extend my thanks to my friends for their support during development and testing, and to my family for their patience and "
        "encouragement throughout the project work.",
    )
    p = doc.add_paragraph()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=16, space_after=0, line_spacing=1.15)
    r = p.add_run(STUDENT_NAME)
    _set_run_font(r, size_pt=12, bold=True)


def _add_abstract(doc: Document):
    _heading_center(doc, "ABSTRACT", size=14, bold=True, uppercase=True, space_before=12, space_after=12)
    _body_para(
        doc,
        f'The project titled "{PROJECT_TITLE}" is an online print shop management platform developed for the printing business "{BUSINESS_NAME}". '
        "The primary objective of the system is to provide customers with a user-friendly interface to browse printing products, place customized "
        "orders, upload designs, and track order status from anywhere. The platform supports a comprehensive product catalog such as visiting cards, "
        "bill books, attendance registers, wedding cards, invoice books, brochures, and other custom printing requirements.",
    )
    _body_para(
        doc,
        "The system provides essential e-commerce functionality including customer registration and login, shopping cart management, secure checkout, "
        "payment integration, and automated order confirmation. An admin dashboard is included for managing products, viewing customer orders, updating "
        "order status, handling uploaded design files, and maintaining inventory records. This approach reduces manual work, improves turnaround time, "
        "and ensures accurate record maintenance.",
    )
    _body_para(
        doc,
        "The proposed implementation can be realized using Python (Flask/Django) for server-side processing, HTML/CSS/JavaScript with Bootstrap for a "
        "responsive interface, and a relational database such as SQLite/MySQL for persistent storage. The system is designed to be scalable, secure, "
        "and efficient, enabling modern digital printing workflows and improved customer experience.",
    )


def _add_toc_like_sample(doc: Document):
    # TOC as a table (like the sample PDF), with deterministic page numbers.
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Chapter\nNo."
    hdr[1].text = "Chapter Titles"
    hdr[2].text = "Page No."
    for i, c in enumerate(hdr):
        for p in c.paragraphs:
            _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15)
            for r in p.runs:
                _set_run_font(r, size_pt=12, bold=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    entries = [
        ("", "Bona-fide Certificate", "ii"),
        ("", "Declaration", "iii"),
        ("", "Acknowledgements", "iv"),
        ("", "Abstract", "v"),
        ("", "Table of Contents", "vi"),
        ("1", "INTRODUCTION", "1"),
        ("2", "OBJECTIVES", "4"),
        ("3", "SYSTEM STUDY", "6"),
        ("4", "EXPERIMENTAL WORK & METHODOLOGY", "12"),
        ("5", "TESTING", "24"),
        ("6", "RESULTS AND DISCUSSION", "30"),
        ("7", "CONCLUSION AND FUTURE WORK", "36"),
        ("8", "REFERENCES", "38"),
        ("9", "APPENDIX", "40"),
    ]
    for ch_no, title, pg in entries:
        row = table.add_row().cells
        row[0].text = ch_no
        row[1].text = title
        row[2].text = pg
        row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for idx, cell in enumerate(row):
            for p in cell.paragraphs:
                _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER if idx != 1 else WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.15)
                for r in p.runs:
                    _set_run_font(r, size_pt=12, bold=(idx != 1 and bool(ch_no)) or (title.isupper()))


def _ensure_font(path_candidates: List[str]) -> Optional[str]:
    for p in path_candidates:
        if os.path.exists(p):
            return p
    return None


def _draw_block_diagram(path: Path, *, title: str, modules: List[str]):
    # Simple academic block diagram (black line boxes) similar to the sample.
    W, H = 1400, 900
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)

    font_path = _ensure_font(
        [
            "C:/Windows/Fonts/times.ttf",
            "C:/Windows/Fonts/timesbd.ttf",
            "C:/Windows/Fonts/timesi.ttf",
            "C:/Windows/Fonts/arial.ttf",
        ]
    )
    font_title = ImageFont.truetype(font_path, 48) if font_path else ImageFont.load_default()
    font_box = ImageFont.truetype(font_path, 30) if font_path else ImageFont.load_default()

    # Top title box
    top_box = (350, 70, 1050, 170)
    d.rectangle(top_box, outline="black", width=4)
    tw = d.textlength(title, font=font_title)
    d.text((700 - tw / 2, 95), title, fill="black", font=font_title)

    # Module boxes
    cols = 3
    rows = math.ceil(len(modules) / cols)
    x0, y0 = 160, 280
    x_gap, y_gap = 380, 160
    box_w, box_h = 320, 90

    centers = []
    for i, m in enumerate(modules):
        r = i // cols
        c = i % cols
        x = x0 + c * x_gap
        y = y0 + r * y_gap
        rect = (x, y, x + box_w, y + box_h)
        d.rectangle(rect, outline="black", width=4)
        mw = d.textlength(m, font=font_box)
        d.text((x + box_w / 2 - mw / 2, y + 27), m, fill="black", font=font_box)
        centers.append(((x + box_w / 2), (y + box_h / 2)))

    # Connect from top box to first row centers
    top_center = ((top_box[0] + top_box[2]) / 2, top_box[3])
    first_row = centers[: min(cols, len(centers))]
    # Branch line
    branch_y = 225
    d.line([top_center, (top_center[0], branch_y)], fill="black", width=4)
    if first_row:
        min_x = min(c[0] for c in first_row)
        max_x = max(c[0] for c in first_row)
        d.line([(min_x, branch_y), (max_x, branch_y)], fill="black", width=4)
        for cx, cy in first_row:
            d.line([(cx, branch_y), (cx, cy - box_h / 2)], fill="black", width=4)
            # arrow head
            d.polygon([(cx - 10, cy - box_h / 2), (cx + 10, cy - box_h / 2), (cx, cy - box_h / 2 + 18)], fill="black")

    # Caption
    caption = "Fig 3.1 Proposed system block diagram"
    cw = d.textlength(caption, font=font_box)
    d.text((W / 2 - cw / 2, H - 90), caption, fill="black", font=font_box)

    img.save(path)


def _draw_dfd_level0(path: Path):
    W, H = 1400, 900
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    font_path = _ensure_font(["C:/Windows/Fonts/times.ttf", "C:/Windows/Fonts/arial.ttf"])
    font = ImageFont.truetype(font_path, 30) if font_path else ImageFont.load_default()
    font_b = ImageFont.truetype(font_path, 34) if font_path else ImageFont.load_default()

    # External entities
    customer = (120, 180, 420, 290)
    admin = (980, 180, 1280, 290)
    d.rectangle(customer, outline="black", width=4)
    d.rectangle(admin, outline="black", width=4)
    d.text((210, 215), "Customer", fill="black", font=font_b)
    d.text((1090, 215), "Admin", fill="black", font=font_b)

    # Process (circle)
    cx, cy, r = 700, 480, 170
    d.ellipse((cx - r, cy - r, cx + r, cy + r), outline="black", width=4)
    d.text((585, 440), "Online Print\nShop System", fill="black", font=font)

    # Data store
    store = (520, 730, 880, 820)
    d.rectangle(store, outline="black", width=4)
    d.text((585, 760), "Database", fill="black", font=font_b)

    # Flows
    def arrow(p1, p2):
        d.line([p1, p2], fill="black", width=4)
        # arrowhead
        ang = math.atan2(p2[1] - p1[1], p2[0] - p1[0])
        ah = 16
        left = (p2[0] - ah * math.cos(ang - 0.5), p2[1] - ah * math.sin(ang - 0.5))
        right = (p2[0] - ah * math.cos(ang + 0.5), p2[1] - ah * math.sin(ang + 0.5))
        d.polygon([p2, left, right], fill="black")

    arrow((420, 235), (530, 380))
    d.text((450, 310), "Order / Login", fill="black", font=font)
    arrow((530, 580), (420, 260))
    d.text((430, 560), "Confirmations", fill="black", font=font)
    arrow((870, 380), (980, 235))
    d.text((840, 310), "Manage Orders", fill="black", font=font)
    arrow((980, 260), (870, 580))
    d.text((890, 560), "Status Updates", fill="black", font=font)
    arrow((700, 650), (700, 730))
    arrow((700, 730), (700, 650))

    caption = "Fig 3.2 Data Flow Diagram (Level-0)"
    cw = d.textlength(caption, font=font)
    d.text((W / 2 - cw / 2, H - 70), caption, fill="black", font=font)

    img.save(path)


def _draw_er_diagram(path: Path):
    W, H = 1400, 900
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    font_path = _ensure_font(["C:/Windows/Fonts/times.ttf", "C:/Windows/Fonts/arial.ttf"])
    font = ImageFont.truetype(font_path, 26) if font_path else ImageFont.load_default()
    font_b = ImageFont.truetype(font_path, 30) if font_path else ImageFont.load_default()

    # Entities as tables
    def entity(x, y, title, fields):
        w = 360
        row_h = 34
        h = 48 + row_h * len(fields)
        d.rectangle((x, y, x + w, y + h), outline="black", width=4)
        d.rectangle((x, y, x + w, y + 48), outline="black", width=4)
        tw = d.textlength(title, font=font_b)
        d.text((x + w / 2 - tw / 2, y + 10), title, fill="black", font=font_b)
        for i, f in enumerate(fields):
            d.text((x + 16, y + 58 + i * row_h), f, fill="black", font=font)
        return (x + w / 2, y + h / 2)

    user_c = entity(120, 130, "USERS", ["userId (PK)", "email", "mobile", "passwordHash", "createdAt"])
    prod_c = entity(520, 130, "PRODUCTS", ["productId (PK)", "name", "basePrice", "category", "imageUrl"])
    order_c = entity(920, 130, "ORDERS", ["orderId (PK)", "userId (FK)", "totalAmount", "paymentMethod", "orderStatus", "createdAt"])
    item_c = entity(520, 520, "ORDER_ITEMS", ["itemId (PK)", "orderId (FK)", "productId (FK)", "qty", "price"])

    # Relationships
    def line(p1, p2, label=None):
        d.line([p1, p2], fill="black", width=4)
        if label:
            lx = (p1[0] + p2[0]) / 2
            ly = (p1[1] + p2[1]) / 2
            d.text((lx - 40, ly - 20), label, fill="black", font=font)

    line((user_c[0] + 160, user_c[1]), (order_c[0] - 160, order_c[1]), "1..N")
    line((order_c[0], order_c[1] + 150), (item_c[0], item_c[1] - 170), "1..N")
    line((prod_c[0], prod_c[1] + 150), (item_c[0], item_c[1] - 170), "1..N")

    caption = "Fig 4.1 Entity Relationship Diagram (ERD)"
    cw = d.textlength(caption, font=font_b)
    d.text((W / 2 - cw / 2, H - 70), caption, fill="black", font=font_b)
    img.save(path)


def _add_figure(doc: Document, image_path: Path, *, width_inches: float = 5.8):
    p = doc.add_paragraph()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=8)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def _add_chapter_title(doc: Document, ch_no: int, title: str):
    _heading_center(doc, f"CHAPTER {ch_no}", size=14, bold=True, uppercase=True, space_before=8, space_after=2)
    _heading_center(doc, title, size=14, bold=True, uppercase=True, space_before=0, space_after=14)


def _add_db_table(doc: Document, table_name: str, rows: List[Tuple[str, str, str, str]]):
    _heading_left(doc, f"Table Name: {table_name}", size=12, bold=True, space_before=12, space_after=6)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "FIELD NAME"
    hdr[1].text = "DATA TYPES"
    hdr[2].text = "WIDTH"
    hdr[3].text = "CONSTRAINTS"
    for c in hdr:
        for p in c.paragraphs:
            _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15)
            for r in p.runs:
                _set_run_font(r, size_pt=12, bold=True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for f, dt, w, cons in rows:
        cells = table.add_row().cells
        cells[0].text = f
        cells[1].text = dt
        cells[2].text = w
        cells[3].text = cons
        for idx, c in enumerate(cells):
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in c.paragraphs:
                _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT if idx in (0, 3) else WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15)
                for r in p.runs:
                    _set_run_font(r, size_pt=12)


def _add_test_table(doc: Document, title: str, cases: List[Tuple[str, str, str, str, str]]):
    _heading_left(doc, title, size=12, bold=True, space_before=12, space_after=6)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Test Case ID"
    hdr[1].text = "Description"
    hdr[2].text = "Action"
    hdr[3].text = "Test Data"
    hdr[4].text = "Expected Result"
    for c in hdr:
        for p in c.paragraphs:
            _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15)
            for r in p.runs:
                _set_run_font(r, size_pt=12, bold=True)

    for row in cases:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
        for c in cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in c.paragraphs:
                _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.15)
                for r in p.runs:
                    _set_run_font(r, size_pt=11)


def _add_code_block(doc: Document, title: str, code: str):
    _heading_left(doc, title, size=12, bold=True, space_before=10, space_after=6)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.clear()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=0, line_spacing=1.15)
    for line in code.rstrip().splitlines():
        run = p.add_run(line + "\n")
        _set_run_font(run, name="Courier New", size_pt=10, bold=False)


def _add_screenshot_placeholder(doc: Document, caption: str):
    _heading_left(doc, caption, size=12, bold=True, space_before=10, space_after=6)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.clear()
    _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0, line_spacing=1.15)
    r = p.add_run("SCREENSHOT PLACEHOLDER")
    _set_run_font(r, size_pt=12, bold=True)
    # Give the placeholder some height.
    cell.height = Inches(3.2)
    # Add a second paragraph for figure caption like sample.
    cp = doc.add_paragraph()
    _set_paragraph(cp, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=8, line_spacing=1.15)
    rr = cp.add_run(caption)
    _set_run_font(rr, size_pt=11, bold=False)


def build():
    doc = Document()
    _doc_defaults(doc)

    # Front matter section: roman numerals starting at i.
    section0 = doc.sections[0]
    _apply_page_setup(section0)
    _set_section_pagenum(section0, start=1, fmt="roman")
    _footer_center_page_number(section0, font_size=11)

    # --- Cover Page (i)
    _add_cover_page(doc)
    doc.add_page_break()

    # --- Bonafide (ii)
    _add_bonafide(doc)
    doc.add_page_break()

    # --- Declaration (iii)
    _add_declaration(doc)
    doc.add_page_break()

    # --- Acknowledgement (iv)
    _add_acknowledgement(doc)
    doc.add_page_break()

    # --- Abstract (v)
    _add_abstract(doc)
    doc.add_page_break()

    # --- Table of Contents (vi)
    _add_toc_like_sample(doc)

    # New section: main matter with arabic numbering starting at 1.
    section_main = doc.add_section(WD_SECTION.NEW_PAGE)
    _apply_page_setup(section_main)
    _set_section_pagenum(section_main, start=1, fmt="decimal")
    _footer_center_page_number(section_main, font_size=11)

    # --- Chapter 1
    _add_chapter_title(doc, 1, "INTRODUCTION")
    _heading_left(doc, "1.1 Overview", size=12, bold=True)
    _body_para(
        doc,
        "Printing press businesses have traditionally relied on walk-in customers, manual job cards, and phone-based order confirmations. "
        "With increasing demand for personalization, rapid turnaround, and transparent pricing, customers expect an online experience comparable "
        "to modern e-commerce platforms. An online print shop enables customers to browse products, upload designs, preview orders, and track delivery "
        "status without visiting the store physically.",
    )
    _body_para(
        doc,
        "The printing industry covers a wide range of products such as business cards, letterheads, brochures, banners, invoices, bill books, "
        "wedding invitations, custom packaging, and branded merchandise. Each product can vary by size, paper type, finishing (lamination, UV, embossing), "
        "quantity, and delivery speed. Therefore, a digital system must support variants, dynamic pricing, design uploads, and reliable order processing workflows.",
    )
    _heading_left(doc, "1.2 Need for Digital Printing Solutions", size=12, bold=True)
    _body_para(
        doc,
        "A digital ordering system reduces manual effort, eliminates repeated communication for order specifications, and improves accuracy by capturing "
        "customer inputs in structured forms. It also provides a consistent experience for repeat customers through saved addresses, order history, and quick reordering. "
        "From a business perspective, it improves operational efficiency by standardizing job intake and reducing errors in the print queue.",
    )
    _body_para(
        doc,
        "For small and medium printing businesses, an online system also acts as a marketing channel. Search engines, social media campaigns, and "
        "business listing platforms can drive customers directly to product pages. Clear product photos, transparent pricing, and fast inquiry handling "
        "increase conversion rates and reduce dependency on local footfall.",
    )
    _heading_left(doc, "1.3 Benefits of Online Ordering", size=12, bold=True)
    _bullets(
        doc,
        [
            "24x7 availability for browsing and ordering without physical visit.",
            "Reduced turnaround time through structured order intake and automated validation.",
            "Higher accuracy of print specifications due to standardized fields and templates.",
            "Improved repeat business through saved designs, reorder options, and order history.",
            "Centralized order data useful for analytics, reporting, and business planning.",
        ],
    )
    _heading_left(doc, "1.3 Purpose of the Project", size=12, bold=True)
    _body_para(
        doc,
        f'The purpose of this project is to design and implement an "{PROJECT_TITLE}" for the print shop business "{BUSINESS_NAME}". '
        "The system supports end-to-end customer ordering including product discovery, cart management, payment selection, design file upload, and order tracking. "
        "It also provides backend capabilities for managing products, orders, payments, and inventory.",
    )
    _heading_left(doc, "1.4 System Users", size=12, bold=True)
    _body_para(
        doc,
        "The system primarily serves two types of users. Customers use the public website to browse products, upload designs, and place orders. "
        "Administrators manage operational workflows such as order validation, status updates, product catalog maintenance, payment reconciliation, and inventory updates. "
        "For security reasons, admin operations are not exposed in the public frontend and are enforced through backend role-based authorization.",
    )
    _heading_left(doc, "1.4 Scope of the Application", size=12, bold=True)
    _bullets(
        doc,
        [
            "Customer-facing website for browsing print products, customizing orders, and placing requests online.",
            "Secure authentication for customers, with options for email/password or mobile OTP-based login.",
            "Shopping cart and checkout modules with address capture and payment method selection.",
            "Order management with status tracking from New to Processing and Completed.",
            "Admin-side capabilities planned as a separate application consuming secured backend APIs.",
        ],
    )
    _heading_left(doc, "1.5 Limitations", size=12, bold=True)
    _bullets(
        doc,
        [
            "Real-time logistics integration is considered future scope; current implementation supports basic delivery status tracking.",
            "Advanced design rendering (server-side PDF proof generation) can be extended using dedicated rendering services.",
            "Inventory forecasting and vendor purchase workflows are basic and can be extended with analytics.",
        ],
    )
    doc.add_page_break()

    # --- Chapter 2
    _add_chapter_title(doc, 2, "OBJECTIVES")
    _body_para(
        doc,
        "The main objectives of the Online Print Shop Management System are as follows:",
        space_after=4,
    )
    _bullets(
        doc,
        [
            "To simplify print order placement through an online product catalog and guided checkout process.",
            "To improve customer accessibility by enabling browsing and ordering from any device at any time.",
            "To reduce manual workload in the print shop by standardizing order intake and capturing structured specifications.",
            "To enable product customization through text inputs and design file uploads.",
            "To maintain accurate records for orders, payments, and customer details for auditability and business analytics.",
            "To support secure authentication and authorization to protect customer data and admin operations.",
            "To streamline inventory tracking for commonly used materials and provide better planning for print jobs.",
        ],
    )
    doc.add_page_break()

    # --- Chapter 3
    _add_chapter_title(doc, 3, "SYSTEM STUDY")
    _heading_left(doc, "3.1 Existing System", size=12, bold=True)
    _body_para(
        doc,
        "In a traditional offline printing press, customers typically visit the shop and provide requirements verbally or through handwritten notes. "
        "The print operator collects details such as product type, quantity, paper quality, and expected delivery date. Design files are often transferred through "
        "messaging apps or physical storage devices. Order tracking is handled using manual registers or spreadsheets, and payment details are maintained separately.",
    )
    _heading_left(doc, "3.2 Disadvantages of Existing System", size=12, bold=True)
    _bullets(
        doc,
        [
            "High dependency on manual communication leading to missed specifications and rework.",
            "Lack of standardized pricing and difficulty in maintaining consistent quotations.",
            "No centralized order history or convenient reorder mechanism for repeat customers.",
            "Limited visibility into order status for customers once the order is placed.",
            "Manual tracking of inventory and print queue leading to delays and operational inefficiency.",
        ],
    )
    _heading_left(doc, "3.3 Proposed System", size=12, bold=True)
    _body_para(
        doc,
        f"The proposed system is a web-based application for {BUSINESS_NAME} that allows customers to browse products, customize orders, upload "
        "design files, and place orders online. The system provides a structured checkout workflow and stores order information in a database. The admin "
        "module is implemented through protected backend APIs so that an admin panel can later be deployed as a separate application.",
    )
    _heading_left(doc, "3.4 Advantages of Proposed System", size=12, bold=True)
    _bullets(
        doc,
        [
            "Improved accessibility and convenience for customers through online ordering.",
            "Faster order management using structured input and automated validations.",
            "Scalability to support more products, categories, and order volume over time.",
            "Better record maintenance for customers, orders, payments, and file uploads.",
            "Enhanced customer experience through order tracking, saved addresses, and reordering.",
        ],
    )

    # Diagrams for Chapter 3
    block_path = ASSETS_DIR / "block_diagram.png"
    if not block_path.exists():
        _draw_block_diagram(
            block_path,
            title="Online Print Shop Management System",
            modules=["Customer", "Admin", "Product Catalog", "Order Management", "Inventory", "Payment Gateway"],
        )
    _add_figure(doc, block_path, width_inches=6.3)
    _heading_left(doc, "3.5 Hardware and Software Requirements", size=12, bold=True)
    _body_para(
        doc,
        "Hardware requirements for the system include a standard computer/server to host the backend and database, and client devices such as laptops or mobile phones "
        "for accessing the web application. For production deployment, a cloud VM or managed platform can be used.",
    )
    _bullets(
        doc,
        [
            "Server: Dual-core CPU or higher, 4GB+ RAM, stable internet connection.",
            "Client: Any modern browser (Chrome/Edge/Firefox/Safari) with JavaScript enabled.",
            "Software: Node.js/Express or Python (Flask/Django), MongoDB/MySQL/SQLite, Git, and a modern code editor.",
        ],
    )
    doc.add_page_break()

    dfd_path = ASSETS_DIR / "dfd_level0.png"
    if not dfd_path.exists():
        _draw_dfd_level0(dfd_path)
    _add_figure(doc, dfd_path, width_inches=6.3)
    _body_para(
        doc,
        "The Level-0 DFD illustrates the interaction between external entities (Customer and Admin) and the central Online Print Shop System. "
        "Customer requests include login, product browsing, order placement, and design upload. Admin actions include managing products and processing orders. "
        "All operational data is persisted in the database.",
    )
    _heading_left(doc, "3.6 Feasibility Study", size=12, bold=True)
    _heading_left(doc, "3.6.1 Technical Feasibility", size=12, bold=True, space_before=8, space_after=4)
    _body_para(
        doc,
        "The system is technically feasible as it uses widely adopted web technologies and standard e-commerce patterns. Hosting, database, and file storage can be handled using "
        "free-tier or low-cost cloud services suitable for a small business initial deployment.",
    )
    _heading_left(doc, "3.6.2 Economic Feasibility", size=12, bold=True, space_before=8, space_after=4)
    _body_para(
        doc,
        "The operational cost of an online ordering system is justified by reduced manual workload, fewer reprints due to error reduction, and higher conversion of online leads. "
        "Using open-source tooling and free-tier services reduces initial investment.",
    )
    # Start the last feasibility subsection on a fresh page to avoid a mostly-empty orphan page.
    doc.add_page_break()
    _heading_left(doc, "3.6.3 Operational Feasibility", size=12, bold=True, space_before=8, space_after=4)
    p_op = _body_para(
        doc,
        "The workflow is designed to match existing print shop processes: collect requirements, validate designs, process print queue, and dispatch delivery. "
        "The system supports incremental adoption where offline customers can still be served while online operations scale up.",
    )
    doc.add_page_break()

    # --- Chapter 4
    _add_chapter_title(doc, 4, "EXPERIMENTAL WORK AND METHODOLOGY")
    _heading_left(doc, "4.1 Proposed System Modules", size=12, bold=True)
    modules = [
        ("4.1.1 User Authentication", "Provides secure customer sign-up and login using email/password and mobile OTP. It validates inputs, issues JWT tokens, and protects secured APIs."),
        ("4.1.2 Product Management", "Maintains the product catalog including categories, pricing, variants, and product images. Supports listing, filtering, and product detail views."),
        ("4.1.3 Cart System", "Allows customers to add products to cart, update quantities deterministically, remove items, and review totals before checkout. Cart state persists using local storage."),
        ("4.1.4 Order Management", "Creates orders from cart items, stores customer delivery details, supports file uploads, and maintains order lifecycle from New to Completed."),
        ("4.1.5 Admin Dashboard (Backend-driven)", "Admin operations are exposed via protected RBAC APIs for viewing orders, updating statuses, managing users and products. The admin UI is designed to be a separate app."),
        ("4.1.6 Payment Integration", "Supports Cash on Delivery and online payments (UPI/Card) through Razorpay Orders API and signature verification."),
        ("4.1.7 Inventory Management", "Tracks stock and consumption of paper, ink, and other printing materials for high-demand products and supports low-stock alerts."),
    ]
    for title, desc in modules:
        _heading_left(doc, title, size=12, bold=True, space_before=10, space_after=4)
        _body_para(doc, desc)

    _heading_left(doc, "4.2 Methodology", size=12, bold=True)
    _heading_left(doc, "4.2.1 Frontend Development", size=12, bold=True, space_before=10, space_after=4)
    _body_para(
        doc,
        "The frontend is designed with responsive layout principles and reusable UI components for product cards, forms, and checkout screens. "
        "Forms are validated in real-time to ensure correct phone numbers, pincodes, and required details. Product pages include image galleries and "
        "support design upload preview for customization.",
    )
    _heading_left(doc, "4.2.2 Backend Development", size=12, bold=True, space_before=10, space_after=4)
    _body_para(
        doc,
        "The backend is implemented using Node.js and Express. It exposes REST APIs for authentication, products, cart checkout, and order management. "
        "JWT-based security is applied for protected routes, and RBAC middleware restricts admin-only endpoints. File uploads are supported using cloud storage "
        "(Cloudinary) or local disk fallback, and the order database is stored in MongoDB for persistent cloud-backed storage.",
    )
    _heading_left(doc, "4.2.3 Database Tables", size=12, bold=True, space_before=10, space_after=4)

    er_path = ASSETS_DIR / "er_diagram.png"
    if not er_path.exists():
        _draw_er_diagram(er_path)
    _add_figure(doc, er_path, width_inches=6.3)

    _add_db_table(
        doc,
        "Users",
        [
            ("id", "ObjectId", "24", "PRIMARY KEY"),
            ("email", "string", "64", "UNIQUE / NULLABLE"),
            ("mobile", "string", "10", "UNIQUE / NULLABLE"),
            ("passwordHash", "string", "255", "NULLABLE"),
            ("provider", "string", "20", "NOT NULL"),
            ("createdAt", "date", "20", "NOT NULL"),
        ],
    )
    doc.add_paragraph()
    _add_db_table(
        doc,
        "Products",
        [
            ("id", "ObjectId", "24", "PRIMARY KEY"),
            ("name", "string", "120", "NOT NULL"),
            ("category", "string", "60", "NOT NULL"),
            ("basePrice", "number", "10", "NOT NULL"),
            ("images", "array", "-", "NOT NULL"),
            ("createdAt", "date", "20", "NOT NULL"),
        ],
    )
    doc.add_page_break()

    _add_db_table(
        doc,
        "Orders",
        [
            ("orderId", "string", "30", "UNIQUE / NOT NULL"),
            ("customerName", "string", "80", "NOT NULL"),
            ("phone", "string", "10", "NOT NULL"),
            ("email", "string", "80", "NULLABLE"),
            ("address", "object", "-", "NOT NULL"),
            ("items", "array", "-", "NOT NULL"),
            ("totalAmount", "number", "12", "NOT NULL"),
            ("paymentMethod", "string", "20", "NOT NULL"),
            ("paymentStatus", "string", "20", "NOT NULL"),
            ("orderStatus", "string", "20", "NOT NULL"),
            ("uploadedFileURL", "string", "255", "NULLABLE"),
            ("createdAt", "date", "20", "NOT NULL"),
        ],
    )
    doc.add_paragraph()
    _add_db_table(
        doc,
        "Payments",
        [
            ("id", "ObjectId", "24", "PRIMARY KEY"),
            ("orderId", "string", "30", "NOT NULL"),
            ("gateway", "string", "20", "NOT NULL"),
            ("gatewayOrderId", "string", "40", "NULLABLE"),
            ("paymentId", "string", "40", "NULLABLE"),
            ("signature", "string", "255", "NULLABLE"),
            ("status", "string", "20", "NOT NULL"),
            ("createdAt", "date", "20", "NOT NULL"),
        ],
    )
    doc.add_page_break()

    _add_db_table(
        doc,
        "Inventory",
        [
            ("id", "ObjectId", "24", "PRIMARY KEY"),
            ("itemName", "string", "120", "NOT NULL"),
            ("unit", "string", "20", "NOT NULL"),
            ("quantityAvailable", "number", "12", "NOT NULL"),
            ("reorderLevel", "number", "12", "NOT NULL"),
            ("updatedAt", "date", "20", "NOT NULL"),
        ],
    )

    doc.add_page_break()

    _heading_left(doc, "4.3 System Architecture", size=12, bold=True)
    _body_para(
        doc,
        "The overall architecture follows a client-server model. The React/Vite frontend communicates with an Express backend through REST APIs. "
        "The backend enforces authentication and RBAC for protected endpoints and stores data in MongoDB. Static product images are served from the public directory, "
        "and design uploads are stored using Cloudinary (or local disk fallback).",
    )
    _bullets(
        doc,
        [
            "Frontend: React + Vite, React Router for page navigation, Tailwind CSS for UI styling.",
            "Backend: Node.js + Express, JWT-based authentication, RBAC middleware for admin routes.",
            "Database: MongoDB Atlas (free tier) with Mongoose ODM schemas.",
            "File Storage: Cloudinary for uploads with cleanup on order deletion; local fallback supported.",
            "Payments: Razorpay order creation and cryptographic signature verification.",
        ],
    )
    _heading_left(doc, "4.4 Order Placement Workflow", size=12, bold=True)
    _body_para(
        doc,
        "The following workflow summarizes how an order is placed and processed in the system:",
        space_after=4,
    )
    for i, step in enumerate(
        [
            "Customer browses products and views product details including images, price, and available variants.",
            "Customer selects quantity and optional customization details (text inputs, notes, design upload).",
            "Item is added to cart; customer reviews cart items, updates quantities, and confirms total amount.",
            "Customer enters delivery details and selects an address (saved address or new address).",
            "Customer selects payment method: Cash on Delivery or Razorpay (UPI/Card).",
            "For Razorpay: backend creates a payment order, Razorpay checkout opens, and payment is processed.",
            "Backend verifies payment signature and stores the final order details with payment status.",
            "Customer is redirected to success page / home page and can track order status from dashboard.",
        ],
        start=1,
    ):
        p = doc.add_paragraph()
        _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4, line_spacing=1.5)
        r = p.add_run(f"{i}. {step}")
        _set_run_font(r, size_pt=12)

    _heading_left(doc, "4.5 API Endpoints (Summary)", size=12, bold=True)
    _body_para(
        doc,
        "The major REST APIs exposed by the backend are summarized below. These APIs are consumed by the public frontend. Admin operations are exposed only via protected /api/admin endpoints.",
    )
    api_table = doc.add_table(rows=1, cols=4)
    api_table.style = "Table Grid"
    hdr = api_table.rows[0].cells
    hdr[0].text = "Method"
    hdr[1].text = "Endpoint"
    hdr[2].text = "Purpose"
    hdr[3].text = "Access"
    for c in hdr:
        for p in c.paragraphs:
            _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15)
            for r in p.runs:
                _set_run_font(r, size_pt=12, bold=True)
    apis = [
        ("POST", "/api/auth/register", "Register with email/password", "Public"),
        ("POST", "/api/auth/login", "Login with email/password", "Public"),
        ("POST", "/api/auth/send-otp", "Send OTP to mobile", "Public (rate-limited)"),
        ("POST", "/api/auth/verify-otp", "Verify OTP and login", "Public"),
        ("GET", "/api/products", "Fetch product catalog", "Public"),
        ("GET", "/api/products/:id", "Fetch product details", "Public"),
        ("POST", "/api/orders", "Create order (COD or paid)", "Authenticated"),
        ("POST", "/api/orders/verify-payment", "Verify Razorpay signature", "Authenticated"),
        ("GET", "/api/user/orders", "Fetch customer order history", "Authenticated"),
        ("GET", "/api/admin/orders", "Admin order list", "ADMIN only"),
    ]
    for m, ep, pur, acc in apis:
        row = api_table.add_row().cells
        row[0].text = m
        row[1].text = ep
        row[2].text = pur
        row[3].text = acc
        for i, cell in enumerate(row):
            for p in cell.paragraphs:
                _set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT if i in (1, 2) else WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.15)
                for r in p.runs:
                    _set_run_font(r, size_pt=11)

    doc.add_page_break()

    # --- Chapter 5
    _add_chapter_title(doc, 5, "TESTING")
    _heading_left(doc, "5.1 Unit Testing", size=12, bold=True)
    _body_para(
        doc,
        "Unit testing validates individual functions and components such as form validators, cart reducer logic, and utility helpers. "
        "This ensures deterministic cart quantity updates (+1/-1), correct input sanitization, and predictable calculation of totals.",
    )
    _add_test_table(
        doc,
        "Unit Testing Test Cases",
        [
            ("UT-01", "Validate phone number input", "Enter alphabets and digits", "phone='98AB12'", "Validation fails with clear error"),
            ("UT-02", "Cart quantity increment", "Click Add to Cart twice", "qty starts at 0", "Quantity becomes 2 deterministically"),
            ("UT-03", "Remove item when qty is 1 and user clicks minus", "Click '-' on qty 1", "qty=1", "Item removed from cart"),
            ("UT-04", "Cart persistence", "Refresh the page", "cart has 3 items", "Items remain in cart from localStorage"),
            ("UT-05", "City autocomplete suggestion filtering", "Type 'Del'", "cityInput='Del'", "Dropdown shows matching cities (e.g., Delhi)"),
        ],
    )
    _heading_left(doc, "5.2 Validation Testing", size=12, bold=True)
    _add_test_table(
        doc,
        "Validation Testing Test Cases",
        [
            ("VT-01", "Full name allows alphabets and spaces only", "Enter invalid symbols", "name='Himanshu#1'", "Error message shown, submission blocked"),
            ("VT-02", "Pincode numeric and 6 digits", "Enter 5 digits", "pincode='56001'", "Error message shown"),
            ("VT-03", "Checkout requires address selection or valid new address", "Submit without address", "address empty", "Place Order disabled / errors displayed"),
            ("VT-04", "OTP length validation", "Enter 4 digits OTP", "otp='1234'", "Error shown, verify blocked"),
            ("VT-05", "Email format validation", "Enter invalid email", "email='abc@'", "Error shown, submission blocked"),
        ],
    )
    doc.add_page_break()

    _heading_left(doc, "5.3 Integration Testing", size=12, bold=True)
    _add_test_table(
        doc,
        "Integration Testing Test Cases",
        [
            ("IT-01", "Create order with COD", "Select COD and place order", "Cart with 2 items", "Order stored in DB with paymentStatus='Pending'"),
            ("IT-02", "Razorpay order creation", "Select UPI/Card and place order", "amount=999", "Backend returns order_id, Razorpay popup opens"),
            ("IT-03", "Payment verification", "Complete Razorpay payment", "valid signature", "Backend verifies signature and marks payment as Paid"),
            ("IT-04", "Cloudinary upload persistence", "Upload PDF design file", "file='logo.pdf'", "Order stores uploadedFileURL and file is accessible"),
            ("IT-05", "RBAC enforcement", "Call admin endpoint with USER token", "role='USER'", "403 Forbidden"),
        ],
    )
    _heading_left(doc, "5.4 Non-Functional Testing", size=12, bold=True)
    _add_test_table(
        doc,
        "Non-Functional Testing Test Cases",
        [
            ("NFT-01", "Performance of product listing", "Load product catalog", "50+ products", "Page loads within acceptable time and images lazy-load"),
            ("NFT-02", "Security of admin APIs", "Call /api/admin/orders without token", "no auth", "403 Forbidden returned"),
            ("NFT-03", "Usability on mobile devices", "Open on mobile viewport", "360x640", "Layout becomes single-column and remains readable"),
            ("NFT-04", "Rate limiting for OTP endpoints", "Send OTP repeatedly", "10 requests/min", "Requests are throttled and error is returned"),
            ("NFT-05", "Data consistency after refresh", "Refresh after cart update", "remove item", "Removed item does not reappear"),
        ],
    )
    doc.add_page_break()

    # --- Chapter 6
    _add_chapter_title(doc, 6, "RESULTS AND DISCUSSION")
    _body_para(
        doc,
        "The Online Print Shop Management System was implemented successfully with a responsive user interface, predictable cart behavior, and a structured checkout flow. "
        "The system improves customer experience by providing a clear product catalog, saving delivery addresses, and supporting both Cash on Delivery and online payment methods.",
    )
    _heading_left(doc, "6.1 Product Ordering Page", size=12, bold=True)
    _body_para(
        doc,
        "The product listing page allows customers to browse printing services and products with images, descriptions, and prices. Each product can be added to cart from both listing and detail pages. "
        "Cart state is centralized to avoid duplication and inconsistent quantities across pages.",
    )
    _heading_left(doc, "6.2 Payment Page", size=12, bold=True)
    _body_para(
        doc,
        "During checkout, customers can choose Cash on Delivery or Razorpay (UPI/Card). For online payments, an order is created on the server and the Razorpay checkout popup is launched. "
        "Payment signatures are verified on the backend to prevent tampering and to ensure secure payment confirmation.",
    )
    _heading_left(doc, "6.3 Order Tracking", size=12, bold=True)
    _body_para(
        doc,
        "After placing an order, customers can view their order history and track order status (New, Processing, Completed, Cancelled). This enables better transparency and reduces support queries. "
        "Admin operations are protected via RBAC so that order updates are accessible only to authorized users.",
    )
    _heading_left(doc, "6.4 Admin Operations (Backend-driven)", size=12, bold=True)
    _body_para(
        doc,
        "The admin workflow is implemented through secured backend APIs. This allows the admin panel to be deployed separately (for example, admin.example.com) without exposing admin routes in the public frontend. "
        "RBAC middleware validates JWT tokens, checks roles, and prevents unauthorized access even if the endpoint URL is known.",
    )
    _heading_left(doc, "6.5 Discussion on Security and Reliability", size=12, bold=True)
    _body_para(
        doc,
        "Security controls include password hashing, OTP expiry, JWT token validation, and rate limiting for authentication endpoints. Payment verification is performed server-side using cryptographic signature checks. "
        "Data persistence in MongoDB ensures that orders and saved addresses remain consistent after refresh, and uploaded design files are retained using cloud storage.",
    )
    # Screenshot placeholders
    _add_screenshot_placeholder(doc, "Fig 6.1 Login Page")
    _add_screenshot_placeholder(doc, "Fig 6.2 Home Page")
    _add_screenshot_placeholder(doc, "Fig 6.3 Product Catalog")
    _add_screenshot_placeholder(doc, "Fig 6.4 Cart and Checkout Page")
    _add_screenshot_placeholder(doc, "Fig 6.5 Delivery Details Form")
    _add_screenshot_placeholder(doc, "Fig 6.6 Payment (Razorpay) Page")
    _add_screenshot_placeholder(doc, "Fig 6.7 Order Success Page")
    _add_screenshot_placeholder(doc, "Fig 6.8 Customer Dashboard / Order History")
    _add_screenshot_placeholder(doc, "Fig 6.9 Product Detail Page")
    _add_screenshot_placeholder(doc, "Fig 6.10 Customization / Design Upload Page")
    doc.add_page_break()

    # --- Chapter 7
    _add_chapter_title(doc, 7, "CONCLUSION AND FUTURE WORK")
    _heading_left(doc, "7.1 Conclusion", size=12, bold=True)
    _body_para(
        doc,
        f'The "{PROJECT_TITLE}" for "{BUSINESS_NAME}" provides a complete solution for managing online print orders. It supports product browsing, '
        "custom order specifications, file uploads, cart and checkout flow, and secure payment options. The system reduces manual effort in order intake, improves "
        "accuracy of specifications, and enhances customer satisfaction through transparent order tracking.",
    )
    _heading_left(doc, "7.2 Future Work", size=12, bold=True)
    _bullets(
        doc,
        [
            "AI-based design recommendations and template suggestions based on product type and industry.",
            "Live delivery tracking with logistics partner integration and SMS/WhatsApp notifications.",
            "Mobile application (Android/iOS) for customers and delivery notifications.",
            "Cloud deployment with autoscaling and centralized monitoring for uptime and performance analytics.",
            "Advanced analytics dashboard for product demand, conversion tracking, and inventory forecasting.",
        ],
    )
    doc.add_page_break()

    # --- Chapter 8
    _add_chapter_title(doc, 8, "REFERENCES")
    refs = [
        "Python Documentation. https://docs.python.org/",
        "Flask Documentation. https://flask.palletsprojects.com/",
        "Bootstrap Documentation. https://getbootstrap.com/docs/",
        "MDN Web Docs (HTML, CSS, JavaScript). https://developer.mozilla.org/",
        "MySQL Documentation. https://dev.mysql.com/doc/",
        "MongoDB Documentation. https://www.mongodb.com/docs/",
        "Razorpay API Documentation. https://razorpay.com/docs/api/",
        "Express.js Documentation. https://expressjs.com/",
    ]
    for i, r in enumerate(refs, start=1):
        _body_para(doc, f"{i}. {r}", space_after=4)
    doc.add_page_break()

    # --- Chapter 9
    _add_chapter_title(doc, 9, "APPENDIX")
    _heading_left(doc, "9.1 Similarity Report (Placeholder)", size=12, bold=True)
    _body_para(doc, "Similarity report will be attached here after plagiarism check.")
    doc.add_page_break()

    _heading_left(doc, "9.2 Source Code (Selected Snippets)", size=12, bold=True)
    _add_code_block(
        doc,
        "Appendix A: Sample Authentication (Flask-style)",
        """
from flask import Flask, request, jsonify
from werkzeug.security import generate_password_hash, check_password_hash

app = Flask(__name__)

@app.post("/auth/register")
def register():
    data = request.json
    email = data.get("email")
    password = data.get("password")
    if not email or not password:
        return jsonify({"message": "Email and password required"}), 400
    password_hash = generate_password_hash(password)
    # Save user in database (pseudo)
    return jsonify({"message": "User registered"}), 201

@app.post("/auth/login")
def login():
    data = request.json
    email = data.get("email")
    password = data.get("password")
    # Fetch user by email and verify
    if not check_password_hash("stored_hash", password):
        return jsonify({"message": "Invalid credentials"}), 401
    return jsonify({"token": "jwt_token_here"}), 200
""",
    )
    _add_code_block(
        doc,
        "Appendix B: Sample Order Placement (Express-style)",
        """
import express from "express";
import { Order } from "../models/Order.js";

const router = express.Router();

router.post("/api/orders", async (req, res) => {
  try {
    const { customerName, phone, address, items, totalAmount, paymentMethod } = req.body;
    if (!customerName || !phone || !address || !items?.length) {
      return res.status(400).json({ message: "Missing required fields" });
    }
    const order = await Order.create({
      customerName,
      phone,
      address,
      items,
      totalAmount,
      paymentMethod,
      paymentStatus: paymentMethod === "COD" ? "Pending" : "Paid",
      orderStatus: "New",
    });
    res.status(201).json({ orderId: order.orderId });
  } catch (err) {
    res.status(500).json({ message: "Failed to create order" });
  }
});

export default router;
""",
    )
    _add_code_block(
        doc,
        "Appendix C: Sample Razorpay Verification (Node.js)",
        """
import crypto from "crypto";

export function verifySignature({ order_id, payment_id, signature }, secret) {
  const body = `${order_id}|${payment_id}`;
  const expected = crypto.createHmac("sha256", secret).update(body).digest("hex");
  return expected === signature;
}
""",
    )
    _add_code_block(
        doc,
        "Appendix D: OTP Send/Verify (Express-style, simplified)",
        """
const otpStore = new Map(); // mobile -> { otp, expiresAt }

export async function sendOtp(req, res) {
  try {
    const { mobile } = req.body;
    if (!/^[0-9]{10}$/.test(mobile)) return res.status(400).json({ message: "Invalid mobile number" });

    const otp = String(Math.floor(100000 + Math.random() * 900000));
    otpStore.set(mobile, { otp, expiresAt: Date.now() + 5 * 60 * 1000 });
    console.log("OTP for", mobile, "=", otp); // Mock SMS
    res.json({ message: "OTP sent" });
  } catch (e) {
    res.status(500).json({ message: "Error sending OTP" });
  }
}

export async function verifyOtp(req, res) {
  const { mobile, otp } = req.body;
  const entry = otpStore.get(mobile);
  if (!entry || Date.now() > entry.expiresAt) return res.status(400).json({ message: "OTP expired" });
  if (otp !== entry.otp) return res.status(400).json({ message: "Invalid OTP" });
  otpStore.delete(mobile);
  // Create or find user and return JWT
  res.json({ token: "jwt_token_here" });
}
""",
    )
    _add_code_block(
        doc,
        "Appendix E: Cart Reducer Logic (React useReducer)",
        """
export function cartReducer(state, action) {
  switch (action.type) {
    case "ADD": {
      const item = state.items[action.product.id];
      const nextQty = (item?.quantity ?? 0) + 1;
      return { ...state, items: { ...state.items, [action.product.id]: { product: action.product, quantity: nextQty } } };
    }
    case "DEC": {
      const item = state.items[action.id];
      if (!item) return state;
      if (item.quantity <= 1) {
        const { [action.id]: _, ...rest } = state.items;
        return { ...state, items: rest };
      }
      return { ...state, items: { ...state.items, [action.id]: { ...item, quantity: item.quantity - 1 } } };
    }
    case "REMOVE": {
      const { [action.id]: _, ...rest } = state.items;
      return { ...state, items: rest };
    }
    default:
      return state;
  }
}
""",
    )
    _add_code_block(
        doc,
        "Appendix F: Product Listing API (Express-style, simplified)",
        """
import express from "express";
import { Product } from "../models/Product.js";

const router = express.Router();

router.get("/api/products", async (req, res) => {
  try {
    const { q, category } = req.query;
    const filter = {};
    if (category) filter.category = category;
    if (q) filter.name = { $regex: q, $options: "i" };
    const products = await Product.find(filter).sort({ createdAt: -1 }).limit(200);
    res.json({ products });
  } catch (e) {
    res.status(500).json({ message: "Failed to load products" });
  }
});

export default router;
""",
    )
    doc.add_page_break()

    _heading_left(doc, "9.3 Screenshots (Placeholders)", size=12, bold=True)
    _add_screenshot_placeholder(doc, "Login Page")
    _add_screenshot_placeholder(doc, "Home Page")
    _add_screenshot_placeholder(doc, "Product Catalog")
    _add_screenshot_placeholder(doc, "Product Detail Page")
    _add_screenshot_placeholder(doc, "Product Customization Page")
    _add_screenshot_placeholder(doc, "Cart Page")
    _add_screenshot_placeholder(doc, "Checkout Page")
    _add_screenshot_placeholder(doc, "Order History")
    _add_screenshot_placeholder(doc, "Admin Dashboard (Separate App)")
    _add_screenshot_placeholder(doc, "Admin Orders Management Page")
    _add_screenshot_placeholder(doc, "Inventory Management Page")

    doc.save(str(OUT_DOCX))
    print(f"Wrote: {OUT_DOCX}")


if __name__ == "__main__":
    build()
